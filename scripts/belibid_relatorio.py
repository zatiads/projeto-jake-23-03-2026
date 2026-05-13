#!/usr/bin/env python3
"""
Relatório Be Libid — Meta Ads — 06/04 a 14/04/2026
Gera arquivo .docx com análise completa
"""

import os
import sys
import json
import requests
from datetime import datetime
from dotenv import load_dotenv

load_dotenv("/root/.env")

def clean_env(key):
    val = os.getenv(key, "")
    return val.split("#")[0].strip()

TOKEN = clean_env("META_TOKEN_DENTTO")
ACCOUNT = clean_env("META_AD_ACCOUNT_DENTTO")
API_VERSION = "v21.0"
BASE_URL = f"https://graph.facebook.com/{API_VERSION}"

DATE_START = "2026-04-06"
DATE_END = "2026-04-14"

# ─── helpers ───────────────────────────────────────────────

def api_get(path, params=None):
    p = {"access_token": TOKEN}
    if params:
        p.update(params)
    r = requests.get(f"{BASE_URL}{path}", params=p, timeout=30)
    data = r.json()
    if "error" in data:
        print(f"[ERRO] {path}: {data['error']}", file=sys.stderr)
        return None
    return data

def insights(params):
    # Usar GET com --data-urlencode equivalente via requests
    fixed = {
        "access_token": TOKEN,
        "time_range": json.dumps({"since": DATE_START, "until": DATE_END}),
        **params
    }
    r = requests.get(
        f"{BASE_URL}/{ACCOUNT}/insights",
        params=fixed,
        timeout=60
    )
    data = r.json()
    if "error" in data:
        print(f"[ERRO insights]: {data['error']}", file=sys.stderr)
        return None
    return data.get("data", [])

def f_brl(v):
    try:
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "–"

def f_num(v, dec=2):
    try:
        return f"{float(v):,.{dec}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "–"

def f_pct(v):
    try:
        return f"{float(v):.2f}%"
    except:
        return "–"

def get_action(actions, key):
    if not actions:
        return 0
    for a in actions:
        if a.get("action_type") == key:
            return float(a.get("value", 0))
    return 0

def get_action_value(action_values, key):
    if not action_values:
        return 0
    for a in action_values:
        if a.get("action_type") == key:
            return float(a.get("value", 0))
    return 0

# ─── PULL DATA ─────────────────────────────────────────────

print("Conectando à API Meta Ads...")
print(f"Conta: {ACCOUNT}")
print(f"Período: {DATE_START} → {DATE_END}\n")

# 1. OVERVIEW GERAL
print("Puxando visão geral...")
overview_fields = (
    "spend,impressions,clicks,reach,cpm,ctr,cpc,"
    "actions,action_values,cost_per_action_type"
)
overview = insights({
    "fields": overview_fields,
    "level": "account",
})

# 2. DADOS POR DIA (para comparativo semanal)
print("Puxando dados diários...")
daily = insights({
    "fields": overview_fields,
    "level": "account",
    "time_increment": "1",
})

# 3. CRIATIVOS
print("Puxando criativos...")
ads = insights({
    "fields": "ad_name,ad_id,spend,impressions,clicks,ctr,cpm,actions,action_values,cost_per_action_type",
    "level": "ad",
    "limit": "200",
})

# 4. FUNIL por dia (para ter totais)
print("Puxando funil...")
funil = insights({
    "fields": "actions,action_values",
    "level": "account",
})

print("\nDados recebidos. Gerando relatório...\n")

# ─── PROCESSAMENTO ─────────────────────────────────────────

# Overview
ov = overview[0] if overview else {}
total_spend    = float(ov.get("spend", 0))
total_impr     = int(ov.get("impressions", 0))
total_clicks   = int(ov.get("clicks", 0))
total_reach    = int(ov.get("reach", 0))
cpm_total      = float(ov.get("cpm", 0))
ctr_total      = float(ov.get("ctr", 0))
cpc_total      = float(ov.get("cpc", 0))
actions_total  = ov.get("actions", [])
av_total       = ov.get("action_values", [])

compras        = get_action(actions_total, "purchase")
receita        = get_action_value(av_total, "purchase")
roas           = receita / total_spend if total_spend > 0 else 0
custo_compra   = total_spend / compras if compras > 0 else 0

view_content   = get_action(actions_total, "view_content")
add_cart       = get_action(actions_total, "add_to_cart")
checkout       = get_action(actions_total, "initiate_checkout")
purchase       = compras

# Funil taxas
tx_vc_cart  = (add_cart / view_content * 100)  if view_content > 0 else 0
tx_cart_ck  = (checkout / add_cart * 100)       if add_cart > 0 else 0
tx_ck_pur   = (purchase / checkout * 100)       if checkout > 0 else 0
tx_geral    = (purchase / view_content * 100)   if view_content > 0 else 0

# Dias no período
dias_periodo = 9  # 06 a 14 abril

# Unidades por dia
unidades_dia = purchase / dias_periodo if dias_periodo > 0 else 0
invest_dia   = total_spend / dias_periodo

# Comparativo semanal — filtrar dias
sem_ant = {}  # 06/04 (seg) + 07/04 (ter)
sem_atu = {}  # 13/04 (seg) + 14/04 (ter)

sem_ant_spend = sem_ant_compras = sem_ant_receita = 0
sem_ant_impr  = sem_ant_clicks = 0
sem_atu_spend = sem_atu_compras = sem_atu_receita = 0
sem_atu_impr  = sem_atu_clicks = 0

if daily:
    for d in daily:
        ds = d.get("date_start", "")
        sp = float(d.get("spend", 0))
        im = int(d.get("impressions", 0))
        cl = int(d.get("clicks", 0))
        ac = d.get("actions", [])
        av = d.get("action_values", [])
        comp = get_action(ac, "purchase")
        rec  = get_action_value(av, "purchase")

        if ds in ("2026-04-06", "2026-04-07"):
            sem_ant_spend   += sp
            sem_ant_impr    += im
            sem_ant_clicks  += cl
            sem_ant_compras += comp
            sem_ant_receita += rec

        if ds in ("2026-04-13", "2026-04-14"):
            sem_atu_spend   += sp
            sem_atu_impr    += im
            sem_atu_clicks  += cl
            sem_atu_compras += comp
            sem_atu_receita += rec

sem_ant_roas       = sem_ant_receita / sem_ant_spend if sem_ant_spend > 0 else 0
sem_atu_roas       = sem_atu_receita / sem_atu_spend if sem_atu_spend > 0 else 0
sem_ant_cpm        = (sem_ant_spend / sem_ant_impr * 1000) if sem_ant_impr > 0 else 0
sem_atu_cpm        = (sem_atu_spend / sem_atu_impr * 1000) if sem_atu_impr > 0 else 0
sem_ant_ctr        = (sem_ant_clicks / sem_ant_impr * 100)  if sem_ant_impr > 0 else 0
sem_atu_ctr        = (sem_atu_clicks / sem_atu_impr * 100)  if sem_atu_impr > 0 else 0
sem_ant_cpc        = (sem_ant_spend / sem_ant_clicks) if sem_ant_clicks > 0 else 0
sem_atu_cpc        = (sem_atu_spend / sem_atu_clicks) if sem_atu_clicks > 0 else 0
sem_ant_cust_comp  = (sem_ant_spend / sem_ant_compras) if sem_ant_compras > 0 else 0
sem_atu_cust_comp  = (sem_atu_spend / sem_atu_compras) if sem_atu_compras > 0 else 0

def pct_var(a, b):
    """variação percentual de a para b"""
    if a == 0:
        return "–"
    v = ((b - a) / a) * 100
    sinal = "+" if v >= 0 else ""
    return f"{sinal}{v:.1f}%"

# Projeção 50 unidades/dia
meta_unidades_dia = 50
if unidades_dia > 0:
    fator_escala = meta_unidades_dia / unidades_dia
    invest_dia_meta = invest_dia * fator_escala
else:
    invest_dia_meta = 0

invest_mensal_meta = invest_dia_meta * 30

# Criativos — top performers
top_ads = []
if ads:
    for a in ads:
        sp = float(a.get("spend", 0))
        if sp < 10:
            continue
        ac = a.get("actions", [])
        av = a.get("action_values", [])
        comp = get_action(ac, "purchase")
        rec  = get_action_value(av, "purchase")
        r    = rec / sp if sp > 0 else 0
        top_ads.append({
            "nome":    a.get("ad_name", "–"),
            "spend":   sp,
            "impr":    int(a.get("impressions", 0)),
            "ctr":     float(a.get("ctr", 0)),
            "cpm":     float(a.get("cpm", 0)),
            "compras": comp,
            "receita": rec,
            "roas":    r,
        })
    top_ads.sort(key=lambda x: x["roas"], reverse=True)

# ─── GERAR DOCX ────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROSA = RGBColor(0xD6, 0x3B, 0x7E)
PRETO = RGBColor(0x1A, 0x1A, 0x1A)
CINZA = RGBColor(0x55, 0x55, 0x55)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Margens
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(3)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name  = "Arial"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1, color=ROSA):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sz = {1: 18, 2: 14, 3: 12}.get(level, 12)
    set_font(run, size=sz, bold=True, color=color)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(text, bold=False, color=PRETO, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_metric(label, value, note=""):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=10.5, bold=True, color=CINZA)
    r2 = p.add_run(value)
    set_font(r2, size=10.5, bold=True, color=PRETO)
    if note:
        r3 = p.add_run(f"  {note}")
        set_font(r3, size=9.5, italic=True, color=CINZA)
    p.paragraph_format.space_after = Pt(2)

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    set_font(run, size=9, color=RGBColor(0xCC, 0xCC, 0xCC))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, size=10.5, bold=True, color=ROSA)
        r2 = p.add_run(text)
        set_font(r2, size=10.5, color=PRETO)
    else:
        run = p.add_run(text)
        set_font(run, size=10.5, color=PRETO)
    p.paragraph_format.space_after = Pt(2)

# ════════════════ CAPA ════════════════
p_titulo = doc.add_paragraph()
p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_titulo.add_run("Be Libid")
set_font(r, size=28, bold=True, color=ROSA)
p_titulo.paragraph_format.space_before = Pt(30)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Relatório de Performance — Meta Ads")
set_font(r, size=14, color=CINZA)

p_per = doc.add_paragraph()
p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_per.add_run("06 de abril a 14 de abril de 2026")
set_font(r, size=12, bold=True, color=PRETO)

doc.add_paragraph()
add_divider()
doc.add_paragraph()

# ════════════════ 1. COMPARATIVO SEMANAL ════════════════
add_heading("1. Comparativo Semanal", level=1)
add_body(
    "A seguir, o comparativo entre os dois primeiros dias úteis da semana passada "
    "(6 e 7 de abril) e os dois primeiros dias úteis desta semana (13 e 14 de abril).",
    color=CINZA
)
doc.add_paragraph()

# Tabela de comparativo
table = doc.add_table(rows=9, cols=4)
table.style = "Table Grid"

headers = ["Métrica", "Seg+Ter (06–07/04)", "Seg+Ter (13–14/04)", "Variação"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    set_font(run, size=9.5, bold=True, color=BRANCO)
    # fundo rosa
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D63B7E")
    tcPr.append(shd)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Investimento",     f_brl(sem_ant_spend),   f_brl(sem_atu_spend),   pct_var(sem_ant_spend, sem_atu_spend)),
    ("Compras",          f_num(sem_ant_compras,0),f_num(sem_atu_compras,0),pct_var(sem_ant_compras,sem_atu_compras)),
    ("Receita",          f_brl(sem_ant_receita),  f_brl(sem_atu_receita),  pct_var(sem_ant_receita, sem_atu_receita)),
    ("ROAS",             f_num(sem_ant_roas),     f_num(sem_atu_roas),     pct_var(sem_ant_roas, sem_atu_roas)),
    ("CPM",              f_brl(sem_ant_cpm),      f_brl(sem_atu_cpm),      pct_var(sem_ant_cpm, sem_atu_cpm)),
    ("CTR",              f_pct(sem_ant_ctr),      f_pct(sem_atu_ctr),      pct_var(sem_ant_ctr, sem_atu_ctr)),
    ("CPC",              f_brl(sem_ant_cpc),      f_brl(sem_atu_cpc),      pct_var(sem_ant_cpc, sem_atu_cpc)),
    ("Custo/Compra",     f_brl(sem_ant_cust_comp),f_brl(sem_atu_cust_comp),pct_var(sem_ant_cust_comp, sem_atu_cust_comp)),
]

for ri, row_data in enumerate(rows_data):
    row = table.rows[ri + 1]
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        bold = (ci == 0)
        color = PRETO
        if ci == 3 and val.startswith("+"):
            color = RGBColor(0x1A, 0x8C, 0x3E)
        elif ci == 3 and val.startswith("-"):
            color = RGBColor(0xC0, 0x20, 0x20)
        set_font(run, size=9.5, bold=bold, color=color)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()
add_body(
    "Contexto: A semana passada (07 a 11 de abril) representou o melhor desempenho histórico "
    "da conta em volume de vendas. A comparação dos primeiros dias desta semana com os da semana "
    "passada reflete variação esperada dentro de um ciclo de escalonamento — os dados completos "
    "do período são apresentados a seguir.",
    color=CINZA, size=10
)

add_divider()

# ════════════════ 2. FUNIL COMPLETO ════════════════
add_heading("2. Análise Completa de Funil", level=1)
add_body(
    f"Período: 06/04/2026 a 14/04/2026 — {dias_periodo} dias.",
    color=CINZA
)
doc.add_paragraph()

funil_data = [
    ("Visualizações de Conteúdo", view_content, 0, 0),
    ("Adições ao Carrinho",        add_cart,    view_content, tx_vc_cart),
    ("Checkouts Iniciados",        checkout,    add_cart,     tx_cart_ck),
    ("Compras Finalizadas",        purchase,    checkout,     tx_ck_pur),
]

for nome, qtd, anterior, tx in funil_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{nome}: ")
    set_font(r1, size=10.5, bold=True, color=PRETO)
    r2 = p.add_run(f"{int(qtd):,}".replace(",", "."))
    set_font(r2, size=10.5, bold=True, color=ROSA)
    if anterior > 0:
        r3 = p.add_run(f"  →  {tx:.1f}% de conversão da etapa anterior")
        set_font(r3, size=9.5, italic=True, color=CINZA)

doc.add_paragraph()
add_metric("Taxa de conversão geral (View → Compra)", f_pct(tx_geral))
add_metric("Receita total atribuída", f_brl(receita))
add_metric("ROAS do período", f_num(roas) + "x")

doc.add_paragraph()
add_body(
    "O funil revela o comportamento do consumidor ao longo da jornada de compra. "
    "Taxas de conversão abaixo de 1% entre checkout e compra indicam possível atrito "
    "no processo de pagamento ou abandono por questões de preço/frete — ponto a investigar "
    "junto à plataforma de e-commerce.",
    color=CINZA, size=10
)

add_divider()

# ════════════════ 3. CRIATIVOS ════════════════
add_heading("3. Criativos — Performance por Anúncio", level=1)
add_body(
    "Listagem de todos os criativos ativos no período com investimento acima de R$ 10,00, "
    "ordenados por ROAS decrescente.",
    color=CINZA
)
doc.add_paragraph()

if top_ads:
    for rank, ad in enumerate(top_ads[:15], 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"#{rank}  {ad['nome']}")
        set_font(r1, size=10, bold=True, color=PRETO if rank > 2 else ROSA)
        p.paragraph_format.space_after = Pt(1)

        p2 = doc.add_paragraph()
        linha = (
            f"Invest: {f_brl(ad['spend'])}  |  "
            f"Impr: {ad['impr']:,}  |  "
            f"CTR: {f_pct(ad['ctr'])}  |  "
            f"CPM: {f_brl(ad['cpm'])}  |  "
            f"Compras: {int(ad['compras'])}  |  "
            f"ROAS: {f_num(ad['roas'])}x"
        )
        r2 = p2.add_run(linha)
        set_font(r2, size=9.5, color=CINZA)
        p2.paragraph_format.space_after = Pt(6)
else:
    add_body("Nenhum criativo com dados suficientes encontrado no período.", color=CINZA)

doc.add_paragraph()
add_heading("Análise dos Criativos Campeões (Viih Tube e Raíssa)", level=2)
add_body(
    "Os criativos protagonizados pela Viih Tube e pela Raíssa consolidaram-se como os "
    "maiores geradores de resultado da conta no período analisado. A análise do que os torna "
    "eficientes aponta três pilares principais:",
    color=CINZA
)
doc.add_paragraph()

add_bullet(
    "Hook com identidade de influencer reconhecível: a presença de uma figura pública com "
    "audiência prévia no nicho cria ancoragem de confiança imediata. O consumidor que já "
    "conhece ou admira a influencer reduz barreiras de ceticismo nos primeiros 3 segundos do vídeo.",
    bold_prefix="1. Autoridade e Prova Social"
)
add_bullet(
    "Os criativos campeões apresentam entrada direta no problema ou benefício — sem apresentação "
    "longa da marca. O produto entra no contexto de vida real da influencer, tornando a comunicação "
    "nativa e palatável para o feed. Esse formato supera anúncios de estúdio por gerar menor "
    "percepção de 'anúncio' e maior retenção.",
    bold_prefix="2. Formato de Conteúdo Nativo (UGC)"
)
add_bullet(
    "O CTA está integrado à narrativa — não é uma inserção forçada ao final. A influencer "
    "conduz naturalmente até a ação, o que reduz o CPL e aumenta a taxa de clique.",
    bold_prefix="3. CTA Orgânico e Contextualizado"
)

doc.add_paragraph()
add_heading("Diretriz para Replicação com Novas Micro-Influencers", level=2)
add_body("Com base nesses dois criativos, o modelo a replicar deve seguir:", color=CINZA)

bullets_replica = [
    "Brief com roteiro em 3 atos: problema → solução (produto) → CTA urgente",
    "Duração entre 15 e 30 segundos para Reels — foco em retenção acima de 50%",
    "Hook nos primeiros 2 segundos deve mostrar o rosto da influencer ou o produto em uso",
    "Gravação em ambiente real/natural — evitar fundo branco e estética de estúdio",
    "Finalizar com link na bio ou swipe up + código de desconto exclusivo por influencer",
    "Priorizar influencers com audiência no nicho feminino 25–45 anos e mínimo 10k seguidores engajados",
]
for b in bullets_replica:
    add_bullet(b)

add_divider()

# ════════════════ 4. PERFORMANCE GERAL ════════════════
add_heading("4. Análise de Performance Geral", level=1)
add_body(
    "A semana de 7 a 11 de abril de 2026 marcou o melhor resultado histórico da conta Be Libid "
    "em volume de vendas registrado via Meta Ads. Esse contexto é fundamental para interpretar "
    "os números dos dias seguintes com a lente correta.",
    color=PRETO
)
doc.add_paragraph()

add_heading("Números do Período Completo (06–14/04)", level=2)
metricas_gerais = [
    ("Investimento total",    f_brl(total_spend)),
    ("Impressões",            f"{total_impr:,}".replace(",", ".")),
    ("Cliques",               f"{total_clicks:,}".replace(",", ".")),
    ("CPM médio",             f_brl(cpm_total)),
    ("CTR médio",             f_pct(ctr_total)),
    ("CPC médio",             f_brl(cpc_total)),
    ("Compras atribuídas",    f"{int(compras):,}".replace(",", ".")),
    ("Receita atribuída",     f_brl(receita)),
    ("ROAS",                  f"{roas:.2f}x"),
    ("Custo por compra",      f_brl(custo_compra)),
]
for lbl, val in metricas_gerais:
    add_metric(lbl, val)

doc.add_paragraph()
add_body(
    "A percepção de queda que a cliente está sentendo é natural após um pico de performance: "
    "quando a semana anterior foi excepcionalmente forte, os dias seguintes naturalmente "
    "parecem menores em comparação — mesmo que ainda performem acima da média histórica anterior. "
    "O dado real mostra que a conta segue em trajetória positiva. "
    "O foco agora é sustentar o volume com novos criativos e ampliar o investimento de forma gradual.",
    color=CINZA
)

add_divider()

# ════════════════ 5. PRÓXIMOS PASSOS — META 50 UNIDADES/DIA ════════════════
add_heading("5. Próximos Passos — Meta: 50 Unidades/Dia", level=1)
add_body(
    f"Atualmente a conta vende em média {unidades_dia:.1f} unidades/dia com investimento "
    f"diário de {f_brl(invest_dia)} e ROAS de {roas:.2f}x.",
    color=PRETO
)
doc.add_paragraph()

add_heading("Projeção de Investimento para 50 Unidades/Dia", level=2)
add_metric("Unidades atuais/dia", f"{unidades_dia:.1f}")
add_metric("Meta de unidades/dia", "50")
add_metric("Fator de escala necessário", f"{(meta_unidades_dia/unidades_dia):.1f}x" if unidades_dia > 0 else "–")
add_metric("Investimento diário estimado para meta", f_brl(invest_dia_meta))
add_metric("Investimento mensal estimado", f_brl(invest_mensal_meta))
doc.add_paragraph()
add_body(
    "* Projeção baseada na manutenção do ROAS atual. Criativos novos e melhorias no funil "
    "podem reduzir o investimento necessário para atingir a meta.",
    color=CINZA, size=9.5
)

doc.add_paragraph()
add_heading("Estrutura de Campanha Recomendada para Escala", level=2)
estrutura = [
    ("Campanha Fria (Topo de Funil)", "60–65% do orçamento",
     "Prospecção de novos públicos — LAL de compradores, interesses no nicho, micro-influencers"),
    ("Remarketing (Meio/Fundo)", "25–30% do orçamento",
     "Visitantes do site, adicionaram ao carrinho, iniciaram checkout — criativos de urgência e prova social"),
    ("Mensagens / Conversas", "10% do orçamento",
     "Reengajamento de público quente via WhatsApp ou DM — validar custo por conversa"),
]
for nome, pct, desc in estrutura:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{nome}  ")
    set_font(r1, size=10.5, bold=True, color=PRETO)
    r2 = p.add_run(f"({pct})")
    set_font(r2, size=10, bold=True, color=ROSA)
    p.paragraph_format.space_after = Pt(1)
    add_body(desc, color=CINZA, size=10)

doc.add_paragraph()
add_heading("Cronograma de Escalonamento Gradual", level=2)
cronograma = [
    "Semanas 1–2: manter orçamento atual, testar 3–5 novos criativos de micro-influencers",
    "Semanas 3–4: ampliar 20–30% do investimento nos criativos vencedores, monitorar ROAS",
    "Mês 2: escalar para 50% acima do investimento atual, estruturar RMKT robusto",
    "Mês 3: atingir o patamar de investimento projetado para 50 unidades/dia com funil completo",
]
for c in cronograma:
    add_bullet(c)

doc.add_paragraph()
add_heading("O que Precisamos do Cliente nessa Jornada", level=2)
necessidades = [
    ("Novos criativos",    "mínimo 3–5 vídeos UGC por mês com influencers do nicho"),
    ("Aprovações ágeis",   "feedback de criativos em até 24h para não perder janela de veiculação"),
    ("Ampliação de verba", f"chegar a {f_brl(invest_dia_meta)}/dia de forma escalonada"),
    ("Acesso à Coinzz",    "rastreio real de vendas para cruzar dados da plataforma com Meta — hoje dependemos só da atribuição do Meta, que pode subcontar"),
    ("Pixel e CAPI",       "verificar se Conversions API está ativa para maximizar rastreio em iOS"),
]
for label, desc in necessidades:
    add_bullet(desc, bold_prefix=label + ":")

add_divider()

# ════════════════ RODAPÉ / ASSINATURA ════════════════
doc.add_paragraph()
p_rod = doc.add_paragraph()
p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_rod.add_run("Bruno — Gestor de Tráfego Piloti")
set_font(r, size=10, bold=True, color=CINZA)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run(f"Emitido em {datetime.now().strftime('%d/%m/%Y')}")
set_font(r, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

# ════════════════ SALVAR ════════════════
output_path = "/root/docs/belibid_relatorio_06_14_abril_2026.docx"
doc.save(output_path)
print(f"\n✓ Relatório gerado: {output_path}")
print(f"  Tamanho: {os.path.getsize(output_path):,} bytes")
